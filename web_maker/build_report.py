from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT_PATH = r"C:\Users\SBS\Desktop\agent02\web_maker\sk_hynix_report.docx"

DARK_BLUE = "1F3864"
LIGHT_BLUE = "D6E4F0"
WHITE = "FFFFFF"


def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)


def set_run_font(run, size_pt=10.5, bold=False, color_rgb=None):
    run.bold = bold
    run.font.name = "맑은 고딕"
    run.font.size = Pt(size_pt)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    if color_rgb:
        run.font.color.rgb = RGBColor(*color_rgb)


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    set_run_font(run)
    return p


def add_heading1(doc, text):
    p = doc.add_heading(text, level=1)
    for run in p.runs:
        run.font.name = "맑은 고딕"
        run.font.size = Pt(14)
        run.bold = True
        run.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    return p


def add_heading2(doc, text):
    p = doc.add_heading(text, level=2)
    for run in p.runs:
        run.font.name = "맑은 고딕"
        run.font.size = Pt(12)
        run.bold = True
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    return p


def build_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"

    # 헤더 행
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_bg(hdr_cells[i], DARK_BLUE)
        p = hdr_cells[i].paragraphs[0]
        p.clear()
        run = p.add_run(h)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.name = "맑은 고딕"
        run.font.size = Pt(10.5)
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 데이터 행
    for idx, row_data in enumerate(rows):
        row = table.add_row()
        bg = LIGHT_BLUE if idx % 2 == 0 else WHITE
        for j, val in enumerate(row_data):
            cell = row.cells[j]
            set_cell_bg(cell, bg)
            p = cell.paragraphs[0]
            p.clear()
            run = p.add_run(val)
            set_run_font(run)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()


def main():
    doc = Document()

    # 페이지 여백
    section = doc.sections[0]
    for attr in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
        setattr(section, attr, Cm(2.54))

    # ── 제목 ──────────────────────────────────────────────
    title_p = doc.add_heading("SK하이닉스 기업 보고서", 0)
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title_p.runs:
        run.font.name = "맑은 고딕"
        run.font.size = Pt(18)
        run.bold = True
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")

    sub_p = doc.add_paragraph("2025-2026 종합 분석")
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub_p.runs[0]
    set_run_font(sub_run, size_pt=13, color_rgb=(0x1F, 0x38, 0x64))
    sub_p.paragraph_format.space_after = Pt(16)

    # ── 섹션 1. 기업 개요 ──────────────────────────────────
    add_heading1(doc, "1. 기업 개요")
    overview_items = [
        "설립: 1983년 현대전자로 설립, 2012년 SK그룹 편입",
        "사명: SK하이닉스(SK hynix Inc.)",
        "본사: 경기도 이천시, 대한민국",
        "상장: 한국거래소(KRX) 000660",
        "사업 분야: DRAM, NAND 플래시, HBM 메모리 반도체 설계·제조",
        "업계 위상: 삼성전자, Micron과 함께 글로벌 3대 메모리 반도체 기업",
    ]
    for item in overview_items:
        add_bullet(doc, item)
    doc.add_paragraph()

    # ── 섹션 2. 주요 제품 ──────────────────────────────────
    add_heading1(doc, "2. 주요 제품")

    add_heading2(doc, "2-1. HBM (High Bandwidth Memory)")
    hbm_items = [
        "HBM3E: 2025년 양산 주력, 12단 적층, 36GB 최고 용량, NVIDIA GB300 탑재",
        "HBM4: 2025년 9월 세계 최초 개발 완료, 동년 4분기 양산 개시",
        "2026년 HBM3E가 전체 HBM 출하량의 약 2/3 차지 전망",
    ]
    for item in hbm_items:
        add_bullet(doc, item)

    add_heading2(doc, "2-2. DRAM")
    dram_items = [
        "DDR5, RDIMM6, MRDIMM7 (1c8 노드 기반)",
        "256GB 3DS RDIMM, LPCAMM2, CXL Memory Module (CMM-DDR5)",
        "AI 서버·PC·모바일 전 플랫폼 대응",
    ]
    for item in dram_items:
        add_bullet(doc, item)

    add_heading2(doc, "2-3. NAND / SSD")
    nand_items = [
        "AI-N 라인업: AI-N P(초고성능), AI-N B(고용량·저가), AI-N D(대용량 데이터셋)",
        "인텔 NAND 사업부 인수(2021) 후 글로벌 2위 NAND 공급사로 도약",
    ]
    for item in nand_items:
        add_bullet(doc, item)
    doc.add_paragraph()

    # ── 섹션 3. 재무 실적 ──────────────────────────────────
    add_heading1(doc, "3. 재무 실적")
    build_table(
        doc,
        headers=["구분", "FY2024", "FY2025", "YoY 성장률"],
        rows=[
            ["매출", "66.19조 원", "97.15조 원", "+46.8%"],
            ["영업이익", "23.47조 원", "47.21조 원", "+101.2%"],
            ["영업이익률", "35%", "49%", "+14%p"],
            ["순이익", "19.80조 원", "42.95조 원", "+116.9%"],
        ],
    )
    finance_notes = [
        "2025년 연간 실적은 창사 이래 최고 기록",
        "HBM 매출이 전년 대비 2배 이상 성장하며 실적 견인",
        "Q4 2024, 삼성전자를 제치고 분기 영업이익 1위 달성",
    ]
    for note in finance_notes:
        add_bullet(doc, note)
    doc.add_paragraph()

    # ── 섹션 4. 시장 지위 ──────────────────────────────────
    add_heading1(doc, "4. 시장 지위")
    build_table(
        doc,
        headers=["제품군", "글로벌 순위", "시장 점유율", "주요 경쟁사"],
        rows=[
            ["DRAM", "1위", "~38.7% (2025 Q2)", "삼성전자(32.7%), Micron"],
            ["NAND", "2위", "~21%", "삼성전자"],
            ["HBM", "1위", "~58% (2026 Q1)", "삼성전자(21%), Micron(21%)"],
        ],
    )
    doc.add_paragraph()

    # ── 섹션 5. 시장 전망 ──────────────────────────────────
    add_heading1(doc, "5. 시장 전망")
    outlook_items = [
        "NVIDIA HBM 주문의 약 70% 공급 중, Vera Rubin 차세대 시스템 주요 공급사 목표",
        "HBM 부족 현상 2028년까지 지속 전망",
        "HBM3E 2026년 가격 약 20% 인상 계획",
        "용인 반도체 클러스터: 2027년 2월 장비 반입, 2030년 상반기 월 36만 장 웨이퍼 생산",
        "M15X(청주) 팹 확장: 2026년 하반기 가동(4만 장/월) → 2027년 8만 장/월",
        "2030~2031년까지 월 DRAM 웨이퍼 생산능력 약 55만 장 → 100만 장 2배 확대 계획",
        "Bank of America: 2026년 글로벌 DRAM 매출 51% 성장, ASP 33% 상승 전망",
    ]
    for item in outlook_items:
        add_bullet(doc, item)

    doc.save(OUTPUT_PATH)
    print(f"저장 완료: {OUTPUT_PATH}")


if __name__ == "__main__":
    main()
