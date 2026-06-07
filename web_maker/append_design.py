from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

FILE_PATH = r"C:\Users\SBS\Desktop\agent02\web_maker\sk_hynix_report.docx"

DARK_BLUE = "1F3864"
LIGHT_BLUE = "D6E4F0"
WHITE = "FFFFFF"
SK_RED   = "E8001C"
SK_ORANGE = "FF6600"


def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)


def set_run_font(run, size_pt=10.5, bold=False, color_rgb=None):
    run.bold = bold
    run.font.name = "맑은 고딕"
    run.font.size = Pt(size_pt)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    if color_rgb:
        run.font.color.rgb = RGBColor(*color_rgb)


def add_heading1(doc, text):
    p = doc.add_heading(text, level=1)
    for run in p.runs:
        run.font.name = "맑은 고딕"
        run.font.size = Pt(14)
        run.bold = True
        run.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    return p


def add_heading2(doc, text):
    p = doc.add_heading(text, level=2)
    for run in p.runs:
        run.font.name = "맑은 고딕"
        run.font.size = Pt(12)
        run.bold = True
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    set_run_font(run)
    return p


def build_color_table(doc):
    """브랜드 컬러 테이블 (컬러 스와치 포함)"""
    headers = ["컬러명", "용도 및 의미", "HEX 코드"]
    rows = [
        ("SK Red", "에너지·열정·기술력 상징. 로고 'SK' 워드마크 적용", "#E8001C"),
        ("SK Orange", "창의성·역동성·따뜻함 상징. 로고 'hynix' 워드마크 적용", "#FF6600"),
        ("White", "순수함·혁신·미래 지향. 행복날개 내부 곡선 라인", "#FFFFFF"),
    ]

    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"

    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_bg(hdr_cells[i], DARK_BLUE)
        p = hdr_cells[i].paragraphs[0]
        p.clear()
        run = p.add_run(h)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.name = "맑은 고딕"
        run.font.size = Pt(10.5)
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    bg_map = [SK_RED, SK_ORANGE, WHITE]
    txt_colors = [
        RGBColor(0xFF, 0xFF, 0xFF),
        RGBColor(0xFF, 0xFF, 0xFF),
        RGBColor(0x22, 0x22, 0x22),
    ]

    for idx, (name, desc, hex_code) in enumerate(rows):
        row = table.add_row()
        bg = LIGHT_BLUE if idx % 2 == 0 else WHITE

        # 컬러명 셀 — 실제 브랜드 컬러로 배경 표시
        set_cell_bg(row.cells[0], bg_map[idx])
        p0 = row.cells[0].paragraphs[0]
        p0.clear()
        r0 = p0.add_run(name)
        r0.bold = True
        r0.font.color.rgb = txt_colors[idx]
        r0.font.name = "맑은 고딕"
        r0.font.size = Pt(10.5)
        r0._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
        p0.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 용도 셀
        set_cell_bg(row.cells[1], bg)
        p1 = row.cells[1].paragraphs[0]
        p1.clear()
        r1 = p1.add_run(desc)
        set_run_font(r1)

        # HEX 셀
        set_cell_bg(row.cells[2], bg)
        p2 = row.cells[2].paragraphs[0]
        p2.clear()
        r2 = p2.add_run(hex_code)
        set_run_font(r2)
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()


def build_design_pillars_table(doc):
    """디자인 철학 3대 축 테이블"""
    headers = ["축", "키워드", "내용"]
    rows = [
        ("기술 (Technology)", "Full Stack AI Memory Creator",
         "고객 문제를 함께 해결하고 생태계와 협력해 더 큰 가치를 창출하는 AI 메모리 창조자"),
        ("문화 (Culture)", "One Team Spirit",
         "협업·도전·다양성을 핵심 가치로, 하나의 팀으로서 최고의 성과를 추구"),
        ("지속가능성 (Sustainability)", "Better World",
         "SK의 DBL(Double Bottom Line) 철학 기반, 사회·경제적 가치를 동시에 창출"),
    ]

    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"

    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_bg(hdr_cells[i], DARK_BLUE)
        p = hdr_cells[i].paragraphs[0]
        p.clear()
        run = p.add_run(h)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.name = "맑은 고딕"
        run.font.size = Pt(10.5)
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for idx, (axis, keyword, desc) in enumerate(rows):
        row = table.add_row()
        bg = LIGHT_BLUE if idx % 2 == 0 else WHITE
        data = [axis, keyword, desc]
        for j, val in enumerate(data):
            set_cell_bg(row.cells[j], bg)
            p = row.cells[j].paragraphs[0]
            p.clear()
            run = p.add_run(val)
            bold = (j < 2)
            set_run_font(run, bold=bold)
            if j < 2:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()


def main():
    doc = Document(FILE_PATH)

    # ── 섹션 6. 브랜드 & 디자인 컨셉 ──────────────────────
    add_heading1(doc, "6. 브랜드 & 디자인 컨셉")

    # 6-1. CI 심볼: 행복날개
    add_heading2(doc, "6-1. CI 심볼: 행복날개 (Wings of Happiness)")
    wings_items = [
        "나비(Butterfly) 형태의 좌우 대칭 심볼 — 반도체 칩이 기기 간 '비행'하듯 데이터를 전달하는 속도와 신뢰를 상징",
        "내부 흰색 곡선 라인이 'S'와 'K'를 형상화 — SK그룹 공통 아이덴티티 계승",
        "두 날개는 SK의 DBL(Double Bottom Line) 철학의 두 기둥, 즉 경제적 가치와 사회적 가치를 동시에 의미",
        "행복날개는 SK그룹 전 계열사 공통 적용 — SK하이닉스는 여기에 반도체 기술 혁신 이미지를 더해 차별화",
    ]
    for item in wings_items:
        add_bullet(doc, item)
    doc.add_paragraph()

    # 6-2. 브랜드 컬러
    add_heading2(doc, "6-2. 브랜드 컬러")
    build_color_table(doc)

    # 6-3. 타이포그래피 & 워드마크
    add_heading2(doc, "6-3. 타이포그래피 & 워드마크")
    typo_items = [
        "워드마크 'SK': SK Red 적용 — 기술력·에너지·주도권을 표현",
        "워드마크 'hynix': SK Orange 적용 — 창의성·역동성·개방성을 표현",
        "서체: 부드럽고 둥근 획의 산세리프(Sans-serif) — 첨단 기술 기업임에도 친근하고 접근하기 쉬운 이미지 전달",
        "대소문자 혼용('SK hynix') — 글로벌 브랜드로서의 독자성 확보",
    ]
    for item in typo_items:
        add_bullet(doc, item)
    doc.add_paragraph()

    # 6-4. 디자인 철학 3대 축
    add_heading2(doc, "6-4. 디자인 철학 3대 축")
    build_design_pillars_table(doc)

    # 6-5. 최신 브랜드 활동
    add_heading2(doc, "6-5. 최신 브랜드 활동 (2024~2025)")
    activity_items = [
        "AI 사명 로고플레이: 'SK hynix' 7개 알파벳(S~X)에 기술·문화·지속가능성을 생성형 AI로 시각화",
        "AI MEMORY SHOW: 모델·배경·제품·음악 전 요소를 AI로 생성 — 브랜드 메시지와 기술력을 통합 표현",
        "반도체 굿즈 팝업스토어: 반도체 회로를 모티브로 한 생활용품 출시, 일반 소비자와 브랜드 접점 확대",
        "비전 슬로건 변경: 'AI Memory Provider' → 'Full Stack AI Memory Creator' (2025 SK AI Summit 발표)",
    ]
    for item in activity_items:
        add_bullet(doc, item)

    doc.save(FILE_PATH)
    print(f"섹션 6 추가 완료: {FILE_PATH}")


if __name__ == "__main__":
    main()
