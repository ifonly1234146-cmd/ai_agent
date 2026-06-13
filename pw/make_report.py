from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)

def set_font(run, name="맑은 고딕", size=None, bold=False, color=None):
    run.font.name = name
    run._r.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size:
        run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)

doc = Document()

# 페이지 여백
section = doc.sections[0]
section.page_width = Cm(21)
section.page_height = Cm(29.7)
section.top_margin = Cm(2.5)
section.bottom_margin = Cm(2.5)
section.left_margin = Cm(3)
section.right_margin = Cm(3)

# ── 제목 ──────────────────────────────
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_r = title_p.add_run("인공지능(AI) 위키백과 요약 보고서")
set_font(title_r, size=20, bold=True, color=(0x1F, 0x49, 0x7D))

# 출처
src_p = doc.add_paragraph()
src_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
src_r = src_p.add_run("출처: 위키백과 한국어판 — ko.wikipedia.org/wiki/인공지능")
set_font(src_r, size=9, color=(0x70, 0x70, 0x70))

date_p = doc.add_paragraph()
date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
date_r = date_p.add_run("작성일: 2026년 06월 13일")
set_font(date_r, size=9, color=(0x70, 0x70, 0x70))

doc.add_paragraph()

# ── 1. 개요 ───────────────────────────
h1 = doc.add_heading("1. 개요", level=1)
h1.runs[0].font.name = "맑은 고딕"
h1.runs[0]._r.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")

intro_texts = [
    "인공지능(人工智能, Artificial Intelligence, AI)은 인간의 학습능력, 추론능력, 지각능력을 인공적으로 구현하려는 컴퓨터 과학의 세부 분야 중 하나이다. 정보공학 분야에 있어 하나의 인프라 기술이기도 하며, 인간을 포함한 동물이 갖고 있는 자연 지능(natural intelligence)과는 구별되는 개념이다.",
    "인공지능은 인간의 지능을 모방한 기능을 갖춘 컴퓨터 시스템으로, 의사 결정, 문제 해결, 학습 등 사람의 인지 능력과 유사한 방식으로 작동한다. 초기 정의는 다트머스 회의(1956)에서 존 매카시(John McCarthy)가 제안한 '기계를 인간 행동의 지식에서와 같이 행동하게 만드는 것'이다.",
]
for text in intro_texts:
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(6)
    for run in p.runs:
        set_font(run, size=10.5)

doc.add_paragraph()

# ── 2. 분류 ───────────────────────────
h2 = doc.add_heading("2. 인공지능의 분류", level=1)
h2.runs[0].font.name = "맑은 고딕"
h2.runs[0]._r.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")

classify_data = [
    ("강인공지능 (Strong AI / AGI)", "인간과 동등하거나 그 이상의 지능을 가진 AI. 범용 문제 해결, 자의식 및 자율적 사고가 가능한 이론적 형태. 아직 실현되지 않음."),
    ("약인공지능 (Weak AI / Narrow AI)", "특정 작업에 특화된 AI. 현재 상용화된 대부분의 AI가 해당. 음성 인식, 이미지 분류, 번역, 추천 시스템 등이 대표적 사례."),
    ("초지능 (Superintelligence)", "인간의 지능을 모든 분야에서 크게 능가하는 가상의 AI. 철학적·윤리적 논쟁의 대상이며 미래 연구 과제로 분류됨."),
]

table = doc.add_table(rows=1, cols=2)
table.style = "Table Grid"
hdr_cells = table.rows[0].cells
hdr_cells[0].text = "구분"
hdr_cells[1].text = "설명"
for cell in hdr_cells:
    set_cell_bg(cell, "1F497D")
    for p in cell.paragraphs:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            set_font(run, size=10, bold=True, color=(0xFF, 0xFF, 0xFF))

for category, desc in classify_data:
    row = table.add_row().cells
    row[0].text = category
    row[1].text = desc
    for run in row[0].paragraphs[0].runs:
        set_font(run, size=10, bold=True)
    for run in row[1].paragraphs[0].runs:
        set_font(run, size=10)

doc.add_paragraph()

# ── 3. 역사 ───────────────────────────
h3 = doc.add_heading("3. 역사 연표", level=1)
h3.runs[0].font.name = "맑은 고딕"
h3.runs[0]._r.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")

history = [
    ("1943–1956", "인공지능의 탄생", "맥컬록·피츠의 신경망 모델(1943), 튜링 테스트 제안(1950), 다트머스 회의에서 'Artificial Intelligence' 용어 공식화(1956)"),
    ("1956–1974", "황금기", "초기 AI 프로그램 개발 붐. 자연어 처리, 문제 해결 알고리즘 연구 활발. 낙관적 전망 팽배."),
    ("1974–1980", "첫 번째 암흑기", "계산 한계와 과도한 기대 붕괴. 연구 자금 대폭 삭감. 'AI 겨울(AI Winter)' 시작."),
    ("1980–1987", "AI 붐 (전문가 시스템)", "전문가 시스템(Expert System) 상용화. 산업 현장 도입 증가. 일본 5세대 컴퓨터 프로젝트 가동."),
    ("1987–1993", "두 번째 암흑기", "전문가 시스템의 한계 노출, 유지보수 비용 급증. 재차 연구비 삭감."),
    ("1993–현재", "현대 AI 르네상스", "머신러닝·딥러닝 발전, GPU 병렬 연산, 빅데이터 등장. ChatGPT 등 대형 언어 모델(LLM) 상용화."),
]

htable = doc.add_table(rows=1, cols=3)
htable.style = "Table Grid"
h_hdr = htable.rows[0].cells
for i, txt in enumerate(["시기", "시대 구분", "주요 내용"]):
    h_hdr[i].text = txt
    set_cell_bg(h_hdr[i], "2E75B6")
    for p in h_hdr[i].paragraphs:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            set_font(run, size=10, bold=True, color=(0xFF, 0xFF, 0xFF))

alt = False
for period, era, desc in history:
    row = htable.add_row().cells
    row[0].text = period
    row[1].text = era
    row[2].text = desc
    if alt:
        for cell in row:
            set_cell_bg(cell, "DEEAF1")
    for run in row[0].paragraphs[0].runs:
        set_font(run, size=9.5, bold=True)
    for run in row[1].paragraphs[0].runs:
        set_font(run, size=9.5, bold=True)
    for run in row[2].paragraphs[0].runs:
        set_font(run, size=9.5)
    alt = not alt

doc.add_paragraph()

# ── 4. 주요 쟁점 ──────────────────────
h4 = doc.add_heading("4. 주요 사회적 쟁점", level=1)
h4.runs[0].font.name = "맑은 고딕"
h4.runs[0]._r.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")

issues = [
    ("거짓정보 및 가짜뉴스", "AI 생성 콘텐츠의 진위 구별 어려움. 딥페이크, 허위 정보 자동 생산 우려."),
    ("일자리 감소", "자동화로 인한 단순·반복 직군 대체 가속화. 직업 구조 재편 필요."),
    ("윤리 문제", "알고리즘 편향, 차별적 의사결정, 자율 무기 사용 등 윤리 기준 수립 시급."),
    ("개인정보 유출", "AI 학습에 사용되는 대규모 데이터의 개인정보 포함 위험성."),
    ("인간 통제력 약화", "고도화된 AI 시스템이 인간의 개입 없이 결정을 내리는 상황 증가."),
    ("사고력 저하", "AI 의존도 증가로 인한 인간 고유의 창의적·비판적 사고 약화 우려."),
]

for issue, desc in issues:
    p = doc.add_paragraph(style="List Bullet")
    r1 = p.add_run(f"{issue}: ")
    set_font(r1, size=10.5, bold=True)
    r2 = p.add_run(desc)
    set_font(r2, size=10.5)

doc.add_paragraph()

# ── 5. 주요 응용 분야 ─────────────────
h5 = doc.add_heading("5. 주요 응용 분야", level=1)
h5.runs[0].font.name = "맑은 고딕"
h5.runs[0]._r.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")

applications = [
    "자연어 처리(NLP): 번역, 챗봇, 문서 요약, 감성 분석",
    "컴퓨터 비전: 이미지·영상 인식, 의료 영상 진단, 자율주행",
    "추천 시스템: 콘텐츠 추천(넷플릭스, 유튜브), 전자상거래",
    "로보틱스: 산업용 로봇, 서비스 로봇, 드론 제어",
    "의료·바이오: 신약 개발, 유전체 분석, 질병 예측",
    "금융: 이상 거래 탐지, 신용 평가, 알고리즘 트레이딩",
]
for app in applications:
    p = doc.add_paragraph(app, style="List Bullet")
    for run in p.runs:
        set_font(run, size=10.5)

doc.add_paragraph()

# ── 6. 미래 전망 ──────────────────────
h6 = doc.add_heading("6. 미래 전망", level=1)
h6.runs[0].font.name = "맑은 고딕"
h6.runs[0]._r.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")

future_text = (
    "위키백과는 AI의 미래에 대해 초지능(Superintelligence)의 등장 가능성과 이에 따른 위험성을 중점적으로 다루고 있다. "
    "단기적으로는 범용 AI(AGI) 달성을 위한 연구가 가속화되고 있으며, 대형 언어 모델(LLM)의 급속한 발전이 이를 뒷받침한다. "
    "장기적으로는 AI가 과학 연구, 의료, 교육, 환경 문제 해결 등 인류의 핵심 과제에 기여할 것으로 기대되나, "
    "안전성 확보와 국제적 거버넌스 체계 구축이 선결 과제로 꼽힌다."
)
p = doc.add_paragraph(future_text)
p.paragraph_format.space_after = Pt(6)
for run in p.runs:
    set_font(run, size=10.5)

doc.add_paragraph()

# ── 푸터 ──────────────────────────────
footer = doc.sections[0].footer
fp = footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
fr = fp.add_run("위키백과 인공지능 페이지 요약 보고서 | 2026.06.13 | ko.wikipedia.org/wiki/인공지능")
set_font(fr, size=8, color=(0x70, 0x70, 0x70))

output_path = r"C:\Users\SBS\Desktop\agent02\pw\인공지능_요약보고서.docx"
doc.save(output_path)
print(f"저장 완료: {output_path}")
